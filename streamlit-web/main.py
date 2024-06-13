import streamlit as st
from docx import Document
from PIL import Image
import base64
import os
import streamlit.components.v1 as components

# Funktion til at læse og vise indholdet fra en Word-fil
def display_word_content(word_path, start_para, end_para):
    doc = Document(word_path)
    content = ""
    for para_num in range(start_para, end_para + 1):
        if para_num < len(doc.paragraphs):
            content += doc.paragraphs[para_num].text + "\n\n"
    return contentimport streamlit as st
from docx import Document
from PIL import Image
import base64
import os
import streamlit.components.v1 as components
import tempfile

# Funktion til at konvertere en Word-fil til billeder af hver side
def convert_word_to_images(word_path):
    doc = Document(word_path)
    with tempfile.TemporaryDirectory() as tempdir:
        image_paths = []
        for i, paragraph in enumerate(doc.paragraphs):
            img_path = os.path.join(tempdir, f"page_{i}.png")
            img = Image.new('RGB', (800, 600), color=(255, 255, 255))
            d = ImageDraw.Draw(img)
            d.text((10, 10), paragraph.text, fill=(0, 0, 0))
            img.save(img_path)
            image_paths.append(img_path)
        return image_paths

# Angiv Word-filens sti
word_path = "streamlit-web/Opgave.docx"

# Tjek om filen eksisterer
if not os.path.exists(word_path):
    st.error(f"Filen {word_path} blev ikke fundet.")
else:
    # Konverter Word-dokumentet til billeder
    image_paths = convert_word_to_images(word_path)

    # Definer opgaverne og deres paragraffer
    tasks = {
        "Forside": 0,  # Forside
        "Opgave 1: Segmentering og målgruppevalg": 1,  # Opgave 1: side 1
        "Opgave 2: Marketingmix": 2,  # Opgave 2: side 2
        "Opgave 3: Udbud - Konkurrence": 3,  # Opgave 3: side 3
        "Opgave 4: Service og kundebetjening": 4,  # Opgave 4: side 4
        "Opgave 5: Forretningsforståelse": 5,  # Opgave 5: side 5
        "Opgave 6: Behov og købemotiv": 6,  # Opgave 6: side 6
    }

    # Tilføj en separat download-sektion
    tasks["Download Word"] = None  # Ingen paragraffer at vise for download

    # Tilføj "Virtuel Butik" til navigationen
    tasks["Virtuel Butik"] = None

    # Angiv stien til dit baggrundsbillede
    background_image_path = "streamlit-web/bg.png"

    # Læs baggrundsbilledet og konverter det til base64
    with open(background_image_path, "rb") as image_file:
        background_image_bytes = image_file.read()
        background_image_base64 = base64.b64encode(background_image_bytes).decode()

    # CSS til at style forsiden med baggrundsbillede og forbedret tekstlæselighed
    st.markdown(
        f"""
        <style>
        .main {{
            background-image: url('data:image/jpg;base64,{background_image_base64}');
            background-size: cover;
            background-position: center;
            background-repeat: no-repeat;
        }}
        .header {{
            text-align: center;
            color: white;
            font-size: 50px;
            font-weight: bold;
            margin-top: 20%;
            text-shadow: 2px 2px 4px rgba(0, 0, 0, 0.7);
            background: rgba(0, 0, 0, 0.7);
            padding: 20px;
            border-radius: 10px;
        }}
        .subheader {{
            text-align: center;
            color: white;
            font-size: 30px;
            font-weight: bold;
            margin-top: 10%;
            text-shadow: 2px 2px 4px rgba(0, 0, 0, 0.3);
            background: rgba(0, 0, 0, 0.5);
            padding: 20px;
            border-radius: 10px;
        }}
        .description {{
            text-align: left;
            color: white;
            font-size: 24px;
            font-weight: normal;
            margin-top: 20px;
            text-shadow: 2px 2px 4px rgba(0, 0, 0, 0.7);
            background: rgba(0, 0, 0, 0.2);
            padding: 20px;
            border-radius: 10px;
        }}
        .description2 {{
            text-align: left;
            color: white;
            font-size: 18px;
            font-weight: normal;
            margin-top: 10px;
            text-shadow: 2px 2px 4px rgba(0, 0, 0, 0.7);
            background: rgba(0, 0, 0, 0.8);
            padding: 15px;
            border-radius: 0px;
        }}
        .description3 {{
            text-align: center;
            color: white;
            font-size: 18px;
            font-weight: normal;
            margin-top: 10px;
            text-shadow: 2px 2px 4px rgba(0, 0, 0, 0.7);
            background: rgba(0, 0, 0, 0.7);
            padding: 5px;
            border-radius: 0px;
        }}
        .blurred {{
            filter: blur(5px);
        }}
        .task-header {{
            text-align: center;
            color: white;
            font-size: 36px;
            font-weight: bold;
            text-shadow: 2px 2px 4px rgba(0, 0, 0, 0.7);
            background: rgba(0, 0, 0, 0.7);
            padding: 20px;
            border-radius: 10px;
            margin-bottom: 20px;
        }}
        </style>
        """,
        unsafe_allow_html=True
    )

    # Streamlit sidebar til navigation
    st.sidebar.title("Navigér til opgave")
    selected_task = st.sidebar.selectbox("", list(tasks.keys()))
    for i in range(40):
        st.sidebar.write("\n")
    st.sidebar.write("© Andreas Lykke Nielsen | Afsætning 2024")

    # Visning baseret på valgt opgave
    if selected_task == "Forside":
        st.markdown("<div class='header'>Eksamensopgave i Afsætning</div>", unsafe_allow_html=True)
        st.markdown("<div class='description'>Denne applikation præsenterer opgaverne fra Word-filen og giver mulighed for at gennemse dem enkeltvis.</div>", unsafe_allow_html=True)
    elif selected_task == "Download Word":
        st.markdown(f"<div class='subheader'>Vil du læse opgaven?</div>", unsafe_allow_html=True)
        with open(word_path, "rb") as file:
            st.download_button(
                label="Download Word",
                data=file,
                file_name="Opgave.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
    elif selected_task == "Virtuel Butik":
        st.markdown("<div class='subheader'>Virtuel Butik</div>", unsafe_allow_html=True)
        st.markdown(f"<div class='description2'>Interaktiv visning af den virtuelle butik</div>", unsafe_allow_html=True)
        components.iframe("https://my.matterport.com/show/?m=vq47Jte1ucv", height=600)
    else:
        page_num = tasks[selected_task]
        st.markdown(f"<div class='task-header'>{selected_task}</div>", unsafe_allow_html=True)
        st.image(image_paths[page_num], use_column_width=True)

        # Prototype til upload af billeder relateret til opgaven (kun på opgavesider)
        st.markdown(f"<div class='description3'>Vælg et billede til visning</div>", unsafe_allow_html=True)
        uploaded_file = st.file_uploader("", type=['png', 'jpg', 'jpeg'])

        if uploaded_file:
            image = Image.open(uploaded_file)
            image_path = f"streamlit-web/uploads/{uploaded_file.name}"
            os.makedirs(os.path.dirname(image_path), exist_ok=True)
            image.save(image_path)
            st.image(image, caption="Relevant billede")

            # Gem filstien til det uploadede billede
            if "uploaded_images" not in st.session_state:
                st.session_state["uploaded_images"] = []
            st.session_state["uploaded_images"].append(image_path)

        # Vis tidligere uploadede billeder
        if "uploaded_images" in st.session_state:
            for image_path in st.session_state["uploaded_images"]:
                image = Image.open(image_path)
                st.image(image, caption="Tidligere uploadet billede")


# Funktion til at finde antal paragraffer i Word-filen
def get_word_para_count(word_path):
    doc = Document(word_path)
    return len(doc.paragraphs)

# Angiv Word-filens sti
word_path = "streamlit-web/Opgave.docx"

# Tjek om filen eksisterer
if not os.path.exists(word_path):
    st.error(f"Filen {word_path} blev ikke fundet.")
else:
    # Hent antallet af paragraffer i Word-filen
    para_count = get_word_para_count(word_path)

    # Definer opgaverne og deres paragraffer
    tasks = {
        "Forside": (0, 0),  # Forside
        "Opgave 1: Segmentering og målgruppevalg": (1, 10),  # Opgave 1: paragraf 1 til 10
        "Opgave 2: Marketingmix": (11, 20),  # Opgave 2: paragraf 11 til 20
        "Opgave 3: Udbud - Konkurrence": (21, 30),  # Opgave 3: paragraf 21 til 30
        "Opgave 4: Service og kundebetjening": (31, 40),  # Opgave 4: paragraf 31 til 40
        "Opgave 5: Forretningsforståelse": (41, 50),  # Opgave 5: paragraf 41 til 50
        "Opgave 6: Behov og købemotiv": (51, 60),  # Opgave 6: paragraf 51 til 60
    }

    # Tilføj en separat download-sektion
    tasks["Download Word"] = (None, None)  # Ingen paragraffer at vise for download

    # Tilføj "Virtuel Butik" til navigationen
    tasks["Virtuel Butik"] = (None, None)

    # Angiv stien til dit baggrundsbillede
    background_image_path = "streamlit-web/bg.png"

    # Læs baggrundsbilledet og konverter det til base64
    with open(background_image_path, "rb") as image_file:
        background_image_bytes = image_file.read()
        background_image_base64 = base64.b64encode(background_image_bytes).decode()

    # CSS til at style forsiden med baggrundsbillede og forbedret tekstlæselighed
    st.markdown(
        f"""
        <style>
        .main {{
            background-image: url('data:image/jpg;base64,{background_image_base64}');
            background-size: cover;
            background-position: center;
            background-repeat: no-repeat;
        }}
        .header {{
            text-align: center;
            color: white;
            font-size: 50px;
            font-weight: bold;
            margin-top: 20%;
            text-shadow: 2px 2px 4px rgba(0, 0, 0, 0.7);
            background: rgba(0, 0, 0, 0.7);
            padding: 20px;
            border-radius: 10px;
        }}
        .subheader {{
            text-align: center;
            color: white;
            font-size: 30px;
            font-weight: bold;
            margin-top: 10%;
            text-shadow: 2px 2px 4px rgba(0, 0, 0, 0.3);
            background: rgba(0, 0, 0, 0.5);
            padding: 20px;
            border-radius: 10px;
        }}
        .description {{
            text-align: left;
            color: white;
            font-size: 24px;
            font-weight: normal;
            margin-top: 20px;
            text-shadow: 2px 2px 4px rgba(0, 0, 0, 0.7);
            background: rgba(0, 0, 0, 0.2);
            padding: 20px;
            border-radius: 10px;
        }}
        .description2 {{
            text-align: left;
            color: white;
            font-size: 18px;
            font-weight: normal;
            margin-top: 10px;
            text-shadow: 2px 2px 4px rgba(0, 0, 0, 0.7);
            background: rgba(0, 0, 0, 0.8);
            padding: 15px;
            border-radius: 0px;
        }}
        .description3 {{
            text-align: center;
            color: white;
            font-size: 18px;
            font-weight: normal;
            margin-top: 10px;
            text-shadow: 2px 2px 4px rgba(0, 0, 0, 0.7);
            background: rgba(0, 0, 0, 0.7);
            padding: 5px;
            border-radius: 0px;
        }}
        .blurred {{
            filter: blur(5px);
        }}
        .task-header {{
            text-align: center;
            color: white;
            font-size: 36px;
            font-weight: bold;
            text-shadow: 2px 2px 4px rgba(0, 0, 0, 0.7);
            background: rgba(0, 0, 0, 0.7);
            padding: 20px;
            border-radius: 10px;
            margin-bottom: 20px;
        }}
        </style>
        """,
        unsafe_allow_html=True
    )

    # Streamlit sidebar til navigation
    st.sidebar.title("Navigér til opgave")
    selected_task = st.sidebar.selectbox("", list(tasks.keys()))
    for i in range(40):
        st.sidebar.write("\n")
    st.sidebar.write("© Andreas Lykke Nielsen | Afsætning 2024")

    # Visning baseret på valgt opgave
    if selected_task == "Forside":
        st.markdown("<div class='header'>Eksamensopgave i Afsætning</div>", unsafe_allow_html=True)
        st.markdown("<div class='description'>Denne applikation præsenterer opgaverne fra Word-filen og giver mulighed for at gennemse dem enkeltvis.</div>", unsafe_allow_html=True)
    elif selected_task == "Download Word":
        st.markdown(f"<div class='subheader'>Vil du læse opgaven?</div>", unsafe_allow_html=True)
        with open(word_path, "rb") as file:
            st.download_button(
                label="Download Word",
                data=file,
                file_name="Opgave.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
    elif selected_task == "Virtuel Butik":
        st.markdown("<div class='subheader'>Virtuel Butik</div>", unsafe_allow_html=True)
        st.markdown(f"<div class='description2'>Interaktiv visning af den virtuelle butik</div>", unsafe_allow_html=True)
        components.iframe("https://my.matterport.com/show/?m=vq47Jte1ucv", height=600)
    else:
        start_para, end_para = tasks[selected_task]
        st.markdown(f"<div class='task-header'>{selected_task}</div>", unsafe_allow_html=True)
        word_content = display_word_content(word_path, start_para, end_para)
        st.markdown(f"<div class='description2'>{word_content}</div>", unsafe_allow_html=True)

        # Prototype til upload af billeder relateret til opgaven (kun på opgavesider)
        st.markdown(f"<div class='description3'>Vælg et billede til visning</div>", unsafe_allow_html=True)
        uploaded_file = st.file_uploader("", type=['png', 'jpg', 'jpeg'])

        if uploaded_file:
            image = Image.open(uploaded_file)
            image_path = f"streamlit-web/uploads/{uploaded_file.name}"
            os.makedirs(os.path.dirname(image_path), exist_ok=True)
            image.save(image_path)
            st.image(image, caption="Relevant billede")

            # Gem filstien til det uploadede billede
            if "uploaded_images" not in st.session_state:
                st.session_state["uploaded_images"] = []
            st.session_state["uploaded_images"].append(image_path)

        # Vis tidligere uploadede billeder
        if "uploaded_images" in st.session_state:
            for image_path in st.session_state["uploaded_images"]:
                image = Image.open(image_path)
                st.image(image, caption="Tidligere uploadet billede")
